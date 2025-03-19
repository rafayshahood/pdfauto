import re
from docx import Document
import data.shared_data as shared_data
import streamlit as st
# LLMWHISPERER_API_KEY="qp6bhMbY3EmpaQqNU2KObHX8fMixMSp0nGcf8GyNnNc"
# LLMWHISPERER_API_KEY="8S_vXzGY7lwCIi0CrUPEGyPojbZu_sxkr7tjG6DqkLg"
# LLMWHISPERER_API_KEY="Z_OCRxI2NjZC3EdpoB5aAWNbHKiaHzE3qO6ll1Z37Vc"

def sanitize_filename(filename):
    """Removes invalid characters from filenames."""
    return re.sub(r'[\/:*?"<>|]', '', filename)  # Remove special characters

def process_cell(cell, replacement_dict, dynamic_counters):
    """
    Recursively processes a cell:
      - Replaces text in the paragraphs using the provided replacement dictionary.
      - Handles dynamic (list‑based) replacements using dynamic_counters.
      - Recursively processes any nested tables within the cell.
    """
    # Process paragraphs in the current cell.
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            for old_text, new_val in replacement_dict.items():
                if old_text in run.text:
                    if isinstance(new_val, list):
                        if not new_val:  # Skip if the list is empty.
                            continue
                        idx = dynamic_counters.get(old_text, 0)
                        replacement_str = new_val[idx] if idx < len(new_val) else new_val[-1]
                        run.text = run.text.replace(old_text, replacement_str)
                        dynamic_counters[old_text] = idx + 1
                    else:
                        run.text = run.text.replace(old_text, new_val)
    # Recursively process any nested tables in this cell.
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                process_cell(nested_cell, replacement_dict, dynamic_counters)

# --------------------------
# Process document function.
def process_document_full(file_path, newHeader, replacement1, replacement2, 
                          allSafetyMeasures, dm2_value, edemaResults, depressed_value, 
                          iteration_index, action, total_pages, check_vertigo=False, check_f=False, check_r=False, palpitation_check= False):
    """
    Processes the Word document and performs the following tasks:
      1. Header text replacement.
      2. Table cell replacements in the first and second columns using dynamic (list‑based) replacements.
      3. Updates safety measure checkboxes.
      4. Updates the DM II checkbox.
      5. Updates the Edema section checkboxes.
      6. Updates the Depressed checkbox.
    The modified document is saved as "merged_modiftotal37.docx".
    """

    # Load the document.
    doc = Document(file_path)
    
    # 1. Header Text Replacement.
    old_header_text = "MINT HOME HEALTH CARE"
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if old_header_text in run.text:
                    run.text = run.text.replace(old_header_text, newHeader)
    
    # 2. Table Cell Replacements.
    # Setup dynamic counters for list-based replacements.
    dynamic_counters1 = { key: 0 for key, val in replacement1.items() if isinstance(val, list) }
    dynamic_counters2 = { key: 0 for key, val in replacement2.items() if isinstance(val, list) }
    
    # 2. Table Cell Replacements.
    # Process first column (cell index 0)
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[0]
            process_cell(cell, replacement1, dynamic_counters1)
    
    # Process second column (cell index 1), including nested tables.
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[1]
            process_cell(cell, replacement2, dynamic_counters2)
                                
    
    # # Convert the safety measures string into a list of lower-case items
    # measures = [s.strip().lower() for s in allSafetyMeasures.split(",")]
    # Split the safety measures string into a list and trim whitespace.
    measures = [m.strip() for m in allSafetyMeasures.split(",") if m.strip()]
    
    # Convert measures to lowercase for checking.
    lower_measures = [m.lower() for m in measures]
    
    # Determine if "cane" and/or "walker" appear in any measure.
    cane_present = any("cane" in m for m in lower_measures)
    walker_present = any("walker" in m for m in lower_measures)

        # Create a new list that excludes any measure that mentions "cane" or "walker".
    filtered_measures = [m for m in measures if "cane" not in m.lower() and "walker" not in m.lower()]
    
    # If either cane or walker is present, add a combined measure.
    if cane_present and walker_present:
        filtered_measures.append("Cane, walker Precautions")
    elif cane_present:
        filtered_measures.append("Cane")
    elif walker_present:
        filtered_measures.append("Walker")
    
    # Build the replacement string: prefix each measure with "☒" and join them with a space.
    replacement_str = " ".join(["☒" + measure for measure in filtered_measures])
    
    # Define the placeholder text in the document. Adjust it if needed.
    placeholder = "Fall Precautions to be replaced"
    
    # Iterate over all runs in all tables and replace the placeholder if found.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement_str)

    
    # 4. DM II Checkbox Update.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "DM II" in run.text:
                            run.text = run.text.replace("☒", "☐")
                            if dm2_value:
                                run.text = run.text.replace("☐", "☒", 1)
                            else:
                                run.text = run.text.replace("☒", "☐")
    
    # 5. Edema Section Checkboxes Update.
    edema_results_lower = edemaResults.lower()
    edema_mapping = {
        "Pitting": "pitting",
        "Non-pitting": "non-pitting",
        "Pacer": "pacer",
        "1+": "+1",
        "2+": "+2",
        "3+": "+3",
        "4+": "+4",
        "Dependent": "dependent",
        "Pedal R/L": "pedal r/l",
        "Dorsum R/L": "dorsum r/l"
    }
    # Uncheck all Edema checkboxes.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for label in edema_mapping.keys():
                            if label in run.text:
                                pattern = r"(☒|☐)(" + re.escape(label) + r")"
                                run.text = re.sub(pattern, r"☐\2", run.text)
    # Check Edema checkboxes.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for label, expected in edema_mapping.items():
                            if label in run.text and expected in edema_results_lower:
                                pattern = r"(☒|☐)(" + re.escape(label) + r")"
                                run.text = re.sub(pattern, r"☒\2", run.text)
    
    # 6. Depressed Checkbox Update.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "Depressed" in run.text:
                            run.text = run.text.replace("☒", "☐")
                            if depressed_value:
                                run.text = run.text.replace("☐", "☒", 1)
                            else:
                                run.text = run.text.replace("☒", "☐")

        # 7. Extra Replacement on the Last Page (only in cell index 1)
    if iteration_index == (total_pages - 1):
        target_text = "(for next visit): continue to implement plan of care as approved by PMD."
        if action == "Reset":
            replacement_text = "for next visit): evaluate Patient/Pcg regarding possible recertification."
        elif action == "Discharge":
            replacement_text = "for next visit): evaluate Patient/Pcg regarding possible discharge."
        else:
            replacement_text = target_text  # Fallback if needed

        # Iterate only over the second column of each table.
        for table in doc.tables:
            for row in table.rows:
                cell = row.cells[1]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if target_text in run.text:
                            run.text = run.text.replace(target_text, replacement_text)

        # Determine correct checkbox symbol based on user input
    vertigo_checkbox = "☒Vertigo" if check_vertigo else "☐Vertigo"  # Checked or Unchecked "Vertigo"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text

                        # Update the Vertigo checkbox
                        text = text.replace("☒Vertigo", "☐Vertigo").replace("☐Vertigo", vertigo_checkbox)

                        run.text = text  # Apply updated text

    # Palpitations Check
    palpitation_checkbox = "☒Palpitations" if palpitation_check else "☐Palpitations" 

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        # Update the Vertigo checkbox
                        text = text.replace("☒Palpitations", "☐Palpitations").replace("☐Palpitations", palpitation_checkbox)

                        run.text = text  # Apply updated text

        # Determine correct checkbox symbols based on user input
    f_checkbox = "☒ R" if check_f else "☐ R"  # Checked or Unchecked "F"
    r_checkbox = "☒ Repeat" if check_r else "☐ Repeat"  # Checked or Unchecked "R"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text

                        # Update checkboxes dynamically
                        text = text.replace("☒ R", "☐ R").replace("☐ R", f_checkbox)  # Set "F"
                        text = text.replace("☒ Repeat", "☐ Repeat").replace("☐ Repeat", r_checkbox)  # Set "R"

                        run.text = text  # Apply updated text
    
     # 1. Extract Patient Name
    extractedResults = shared_data.data['extraction_results']
    patient_name = extractedResults['patientDetails']["name"]  # Already in "Last, First" format

    # 2. Extract Disease Name for this page
    page_key = f"page{iteration_index+1}"
    disease_name = st.session_state["disease_mapping"].get(page_key, "Unknown Disease")

    if iteration_index == 9:
        disease_name = st.session_state["disease_mapping"].get(page_key, "Discharge")

    # **Sanitize filename to remove invalid characters**
    safe_disease_name = sanitize_filename(disease_name)
    output_file = f"{patient_name}, {safe_disease_name} - {page_key}.docx"
    # Save the modified document with a name based on iteration (e.g., "page1.docx", "page2.docx", etc.)
    # output_file = f"page{iteration_index+1}.docx"

    doc.save(output_file)
    return output_file


