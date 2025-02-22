# get values within range
import random
import shared_data
import re
from docx import Document
import streamlit as st
import json
from datetime import datetime, timedelta

def adjust_dates(appointment_dates, appointment_times, constipation=False):
    """
    Adjust appointment dates based on appointment times and the constipation flag.
    
    - If constipation is True, subtract one day from every appointment date.
    - Otherwise, only subtract one day for appointments whose starting time is before 10:00 AM.
    
    Parameters:
      appointment_dates (list): List of date strings (format: "dd/mm/YYYY").
      appointment_times (list): List of time strings (e.g., "4:00-4:45"). The first time in the range is used.
      constipation (bool): Flag indicating whether constipation is true.
      
    Returns:
      list: Adjusted date strings in "dd/mm/YYYY" format.
    """
    new_dates = []
    for date_str, time_str in zip(appointment_dates, appointment_times):
        dt = datetime.strptime(date_str, "%d/%m/%Y")
        # Extract the first (starting) time from the range, e.g., "4:00-4:45" -> "4:00"
        start_time_str = time_str.split('-')[0].strip() if '-' in time_str else time_str.strip()
        t = datetime.strptime(start_time_str, "%H:%M")
        cutoff = datetime.strptime("10:00", "%H:%M")
        if constipation:
            new_dt = dt - timedelta(days=1)
        else:
            # Only adjust if the starting time is before 10:00 AM.
            new_dt = dt - timedelta(days=1) if t < cutoff else dt
        new_dates.append(new_dt.strftime("%d/%m/%Y"))
    return new_dates
# --------------------------
# Helper functions for generating random values

def get_random_bp(systolic_min=130, systolic_max=145, diastolic_min=65, diastolic_max=87):
    systolic = random.randint(systolic_min, systolic_max)
    diastolic = random.randint(diastolic_min, diastolic_max)
    return f"{systolic}/{diastolic}"

def get_random_value(min_value, max_value, roundTo=1, is_integer=False):
    value = random.uniform(min_value, max_value)
    if is_integer:
        return round(value)
    return round(value, roundTo)

def getRangeValues():
    tempVal = str(get_random_value(97.7, 99.5))
    hrVal = str(get_random_value(60, 100, is_integer=True))
    rrVal = str(get_random_value(16, 20, is_integer=True))
    oxygenVal = str(get_random_value(95, 99, is_integer=True))
    bsFastVal = str(get_random_value(140, 180, is_integer=True))
    bsRapidVal = str(get_random_value(190, 250, is_integer=True))
    random_bp = str(get_random_bp())
    return tempVal, hrVal, rrVal, oxygenVal, bsFastVal, bsRapidVal, random_bp

def getRangeValuesArray(noOfVal):
    # Generate arrays for n iterations.
    tempArray = []
    hrArray = []
    rrArray = []
    oxygenArray = []
    bsFastArray = []
    bsRapidArray = []
    random_bpArray = []

    for i in range(noOfVal):
        tempVal, hrVal, rrVal, oxygenVal, bsFastVal, bsRapidVal, random_bp = getRangeValues()
        tempArray.append(tempVal)
        hrArray.append(hrVal)
        rrArray.append(rrVal)
        oxygenArray.append(oxygenVal)
        bsFastArray.append(bsFastVal)
        bsRapidArray.append(bsRapidVal)
        random_bpArray.append(random_bp)
    return tempArray, hrArray, oxygenArray, bsFastArray, bsRapidArray, random_bpArray 

# --------------------------
# Process document function.
def process_document_full(file_path, newHeader, replacement1, replacement2, 
                          allSafetyMeasures, dm2_value, edemaResults, depressed_value, painScaleArray):
    """
    Processes the Word document and performs the following tasks:
      1. Header text replacement.
      2. Table cell replacements in the first and second columns using dynamic (list‑based) replacements.
      3. Updates safety measure checkboxes.
      4. Updates the DM II checkbox.
      5. Updates the Edema section checkboxes.
      6. Updates the Depressed checkbox.
    The modified document is saved as "merged_modiftotal37.docx".
    """

    # Load the document.
    doc = Document(file_path)
    
    # 1. Header Text Replacement.
    old_header_text = "MINT HOME HEALTH CARE"
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if old_header_text in run.text:
                    run.text = run.text.replace(old_header_text, newHeader)
    
    # 2. Table Cell Replacements.
    # Setup dynamic counters for list-based replacements.
    dynamic_counters1 = { key: 0 for key, val in replacement1.items() if isinstance(val, list) }
    dynamic_counters2 = { key: 0 for key, val in replacement2.items() if isinstance(val, list) }
    
    # Process first column (cell index 0)
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[0]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for old_text, new_val in replacement1.items():
                        if old_text in run.text:
                            if isinstance(new_val, list):
                                if not new_val:  # new_val is empty
                                    continue  # or set replacement_str = "" and/or log a warning
                                idx = dynamic_counters1[old_text]
                                replacement_str = new_val[idx] if idx < len(new_val) else new_val[-1]
                                run.text = run.text.replace(old_text, replacement_str)
                                dynamic_counters1[old_text] += 1
                            else:
                                run.text = run.text.replace(old_text, new_val)
    
    # Process second column (cell index 1)
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[1]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for old_text, new_val in replacement2.items():
                        if old_text in run.text:
                            if isinstance(new_val, list):
                                if not new_val:  # new_val is empty
                                    continue  # or set replacement_str = "" and/or log a warning
                                idx = dynamic_counters2[old_text]
                                replacement_str = new_val[idx] if idx < len(new_val) else new_val[-1]
                                run.text = run.text.replace(old_text, replacement_str)
                                dynamic_counters2[old_text] += 1
                            else:
                                run.text = run.text.replace(old_text, new_val)
    
    # 3. Safety Measures Checkboxes Update.
    measures = [s.strip().lower() for s in allSafetyMeasures.split(",")]
    safety_measures_mapping = {
        "Bleeding Precautions": "bleeding precautions",
        "Fall Precautions": "fall precautions",
        "Clear pathways": "clear pathways",
        "Infection control measures": "infection control",
        "Cane, walker Precautions": ("cane", "walker"),
        "Universal Precautions": "universal precautions",
        "Other:911 protocols": "911 protocol"
    }
    # Uncheck all safety measure checkboxes.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for label in safety_measures_mapping.keys():
                            if label in run.text:
                                pattern = r"(☒|☐)(" + re.escape(label) + r")"
                                run.text = re.sub(pattern, r"☐\2", run.text)
    # Check safety measure checkboxes based on provided measures.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for label, expected in safety_measures_mapping.items():
                            if label in run.text:
                                should_check = False
                                if isinstance(expected, tuple):
                                    for item in expected:
                                        if item in measures:
                                            should_check = True
                                            break
                                else:
                                    if expected in measures:
                                        should_check = True
                                
                                pattern = r"(☒|☐)(" + re.escape(label) + r")"
                                run.text = re.sub(pattern, r"☒\2" if should_check else r"☐\2", run.text)
    
    # 4. DM II Checkbox Update.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "DM II" in run.text:
                            run.text = run.text.replace("☒", "☐")
                            if dm2_value:
                                run.text = run.text.replace("☐", "☒", 1)
                            else:
                                run.text = run.text.replace("☒", "☐")
    
    # 5. Edema Section Checkboxes Update.
    edema_results_lower = edemaResults.lower()
    edema_mapping = {
        "Pitting": "pitting",
        "Non-pitting": "non-pitting",
        "Pacer": "pacer",
        "1+": "+1",
        "2+": "+2",
        "3+": "+3",
        "4+": "+4",
        "Dependent": "dependent",
        "Pedal R/L": "pedal r/l",
        "Dorsum R/L": "dorsum r/l"
    }
    # Uncheck all Edema checkboxes.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for label in edema_mapping.keys():
                            if label in run.text:
                                pattern = r"(☒|☐)(" + re.escape(label) + r")"
                                run.text = re.sub(pattern, r"☐\2", run.text)
    # Check Edema checkboxes.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for label, expected in edema_mapping.items():
                            if label in run.text and expected in edema_results_lower:
                                pattern = r"(☒|☐)(" + re.escape(label) + r")"
                                run.text = re.sub(pattern, r"☒\2", run.text)
    
    # 6. Depressed Checkbox Update.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "Depressed" in run.text:
                            run.text = run.text.replace("☒", "☐")
                            if depressed_value:
                                run.text = run.text.replace("☐", "☒", 1)
                            else:
                                run.text = run.text.replace("☒", "☐")
    
    # Save the modified document.
    output_file = "merged_modiftotal37.docx"
    doc.save(output_file)
    return output_file

# --------------------------
# Example usage:
# (Make sure that extractedResults, mainContResponse, and other external variables are defined.)
def fillDoc():
    # Example external variables (replace with your own logic or data)
    extractedResults = shared_data.data['extraction_results']
    constipation = extractedResults['diagnosis']['constipated']
    sn_name = shared_data.data['sn_name']
    appointment_dates = shared_data.data['appointment_dates']
    appointment_times = shared_data.data['appointment_times']

    # extractedResults = {
    #     'extraDetails': {
    #         'can': True,
    #         'walker': False,
    #         'safetyMeasures': "Fall Precautions, Infection control measures",
    #         'safetyMeasuresCont': "",
    #         'edema': "Pedal R/L, Pitting +1"
    #     },
    #     'patientDetails': {
    #         'providerName': "New Provider Name",
    #         'medicalRecordNo': "1234567890"
    #     },
    #     'diagnosis': {
    #         'painIn': "Lower back pain"
    #     },
    #     'medications': {
    #         'painMedications': "Tylenol"
    #     }
    # }

    # For the dynamic replacement, assume mainContResponse is defined.
    mainContResponse = shared_data.mainContResponse

    # Get Action
    getAction = shared_data.data['action'] 
    if getAction == "Discharge":
        valuesToGet = 9
        wordFileName = './noteTemplates/9-page-version.docx'  # Ensure this file exists in your working directory
    elif getAction == "Reset":
        valuesToGet = 10
        wordFileName = './noteTemplates/10-page-version.docx'  # Ensure this file exists in your working directory

    # Generate arrays for 10 iterations.
    tempArray = []
    hrArray = []
    rrArray = []
    oxygenArray = []
    bsFastArray = []
    bsRapidArray = []
    random_bpArray = []

    tempArray, hrArray, oxygenArray, bsFastArray, bsRapidArray, random_bpArray = getRangeValuesArray(valuesToGet)

    if extractedResults['extraDetails']['can'] and extractedResults['extraDetails']['walker']:
        canWalkerText = 'can,walker'
    elif extractedResults['extraDetails']['can']:
        canWalkerText = 'can'
    elif extractedResults['extraDetails']['walker']:
        canWalkerText = 'walker'
    else:
        canWalkerText = ''

    headerPage = extractedResults['patientDetails']['providerName']
    


    painScaleArray = ['5/10', '4/10', '4/10', '4/10', '4/10', '3/10', '3/10', '3/10', '2/10']
    if getAction == "Discharge":
        painScaleArray.append('4/10')

    replacements_first_col = {
        'cane,walker': canWalkerText,
        'Pain in Lower back, Neck, Joints': extractedResults['diagnosis']['painIn'],
        '4/10': [pain for pain in painScaleArray],  # Dynamic list replacement using painScaleArray.
        'Tylenol 325 mg. 1 tablet by moiuth daily': extractedResults['medications']['painMedications'],
        '05/07/23': adjust_dates(appointment_dates, appointment_times, constipation)
        }

    replacements_second_col = {
        'T- 96.8': ["T- " + temp for temp in tempArray],
        'HR- 66': ["HR- " + hr for hr in hrArray],
        'RR -16': ["RR - " + rr for rr in rrArray],
        'BS 198': ["BS " + bsFast for bsFast in bsFastArray],
        'Sitting 142/89': ["Sitting " + rb for rb in random_bpArray],
        'PARKER, PETER/LVN': sn_name,
        'MR# 022-001': "MR# " + extractedResults['patientDetails']['medicalRecordNo'][-7:],
        'PORK, JOHN': extractedResults['patientDetails']["name"],
        '05/08/2023': ["      " + date for date in appointment_dates],
        '18:00-18:45': ["      " + time for time in appointment_times],
        'new text1': [json.loads(response).get("text1", "") for response in mainContResponse.values()],
        'replacement text': [json.loads(response).get("text2", "") for response in mainContResponse.values()],

    }

    dm2_value = True
    depressed_value = True
    allSafetyMeasures = extractedResults['extraDetails']['safetyMeasures'] + ', ' + extractedResults['extraDetails']['safetyMeasuresCont']
    edemaResults = extractedResults['extraDetails']['edema']

    # Process the document and get the output filename.
    output_file = process_document_full(wordFileName, headerPage, replacements_first_col, replacements_second_col, 
                                        allSafetyMeasures, dm2_value, edemaResults, depressed_value, painScaleArray)

    # Provide a download button using Streamlit.
    with open(output_file, "rb") as doc_file:
        st.download_button(
            label="Download Modified Document",
            data=doc_file,
            file_name=output_file,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
